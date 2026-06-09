"""
utils/word_generator.py
Genera documentos Word (.docx) con python-docx.
  - Informe mensual de actividades
"""

from __future__ import annotations
import io
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TEAL_RGB = RGBColor(0, 105, 92)
GRAY_RGB = RGBColor(107, 114, 128)
WHITE    = RGBColor(255, 255, 255)

MESES = ["","Enero","Febrero","Marzo","Abril","Mayo","Junio",
         "Julio","Agosto","Septiembre","Octubre","Noviembre","Diciembre"]


def _set_cell_bg(cell, hex_color: str):
    """Aplica color de fondo a una celda de tabla."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def generar_word_informe(mes: int, anio: int, stats: dict) -> bytes:
    doc = Document()

    # Márgenes
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Encabezado ────────────────────────────────────
    h = doc.add_heading("", level=0)
    run = h.add_run("INFORME MENSUAL DE ACTIVIDADES")
    run.font.color.rgb = TEAL_RGB
    run.font.size      = Pt(18)
    h.alignment        = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(f"Período: {MESES[mes]} {anio}  ·  "
                    f"Hospital Nacional Ramiro Prialé Prialé – EsSalud")
    r.font.size      = Pt(10)
    r.font.color.rgb = GRAY_RGB

    doc.add_paragraph()

    # ── Tabla KPI ─────────────────────────────────────
    doc.add_heading("Resumen de Actividades", level=2)
    kpi_rows = [
        ("Total de Atenciones",   str(stats.get("atenciones", 0))),
        ("Pacientes Atendidos",   str(stats.get("activos",    0))),
        ("Pacientes HD",          str(stats.get("hd",         0))),
        ("Pacientes DIPAC/ERCA",  str(stats.get("dipac",      0))),
    ]
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    # encabezado
    for i, txt in enumerate(["Indicador", "Valor"]):
        cell = tbl.rows[0].cells[i]
        cell.text = txt
        _set_cell_bg(cell, "00695C")
        cell.paragraphs[0].runs[0].font.color.rgb = WHITE
        cell.paragraphs[0].runs[0].font.bold      = True
    # filas
    for indicador, valor in kpi_rows:
        row = tbl.add_row()
        row.cells[0].text = indicador
        row.cells[1].text = valor
        row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # ── Distribución por tipo ──────────────────────────
    doc.add_heading("Distribución por Tipo de Intervención", level=2)
    tipos  = stats.get("tipos",  [])
    counts = stats.get("counts", [])
    total  = sum(counts) or 1

    tbl2 = doc.add_table(rows=1, cols=3)
    tbl2.style = "Table Grid"
    for i, txt in enumerate(["Tipo de Atención", "N°", "%"]):
        cell = tbl2.rows[0].cells[i]
        cell.text = txt
        _set_cell_bg(cell, "00695C")
        cell.paragraphs[0].runs[0].font.color.rgb = WHITE
        cell.paragraphs[0].runs[0].font.bold      = True

    for t_nom, cnt in zip(tipos, counts):
        row = tbl2.add_row()
        row.cells[0].text = t_nom
        row.cells[1].text = str(cnt)
        row.cells[2].text = f"{cnt/total*100:.1f}%"
        row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Fila total
    row_tot = tbl2.add_row()
    row_tot.cells[0].text = "TOTAL"
    row_tot.cells[1].text = str(total)
    row_tot.cells[2].text = "100%"
    for c in row_tot.cells:
        _set_cell_bg(c, "E0F2F1")
        for p in c.paragraphs:
            for r in p.runs:
                r.font.bold = True
    for c in row_tot.cells[1:]:
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # ── Conclusiones ──────────────────────────────────
    doc.add_heading("Conclusiones y Recomendaciones", level=2)
    conclusiones = (
        f"Durante el mes de {MESES[mes]} {anio}, el servicio de Psiconefrología del "
        f"Hospital Nacional Ramiro Prialé Prialé atendió un total de {stats.get('atenciones',0)} "
        f"pacientes, de los cuales {stats.get('hd',0)} pertenecen al programa de Hemodiálisis y "
        f"{stats.get('dipac',0)} al programa DIPAC/ERCA. Se recomienda continuar con el seguimiento "
        f"psicológico individualizado, enfatizando la adherencia al tratamiento y el soporte "
        f"familiar como factores protectores en la población con enfermedad renal crónica."
    )
    p = doc.add_paragraph(conclusiones)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_paragraph()
    doc.add_paragraph()

    # ── Firma ─────────────────────────────────────────
    firma = doc.add_paragraph("_____________________________")
    firma.add_run("\nPsicólogo/a – Servicio de Psiconefrología\n"
                  "Hospital Nacional Ramiro Prialé Prialé")
    firma.runs[-1].font.color.rgb = GRAY_RGB
    firma.runs[-1].font.size      = Pt(9)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
