from fpdf import FPDF
import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import pandas as pd

class FichaPDF(FPDF):
    def header(self):
        try:
            self.image('logo_essalud.png', 10, 8, 25) 
        except: pass
        self.set_left_margin(40) 
        self.set_font('Arial', 'B', 8) 
        self.set_text_color(100, 100, 100)
        self.cell(0, 5, 'HOSPITAL NACIONAL RAMIRO PRIALE PRIALE ESSALUD - SERVICIO DE PSICOLOGIA (HD - DIPAC)', ln=1, align='L')
        self.set_left_margin(10)
        self.ln(8)

def generar_pdf_ficha(p, datos, grafico_path=None):
    pdf = FichaPDF()
    pdf.add_page()
    epw = pdf.epw 
    
    # --- CONFIGURACIÓN DE COHERENCIA TOTAL ---
    SIZE = 10         
    H_CELDA = 7       
    BORDES = (180, 180, 180) 
    S_ESPACIO = 6     
    
    pdf.set_draw_color(*BORDES)
    pdf.set_text_color(0, 0, 0)
    
    # Título Principal
    pdf.set_font('Arial', 'B', 12)
    pdf.set_x(10)
    pdf.cell(epw, 10, 'FICHA PSICOLÓGICA', ln=1, align='C', border='B')
    pdf.ln(S_ESPACIO)
    
    # =========================================================
    # I. DATOS DE FILIACIÓN (REORGANIZADO)
    # =========================================================
    pdf.set_fill_color(245, 245, 245)
    pdf.cell(epw, H_CELDA, 'I. Datos de Filiación y Estado Clínico', ln=1, fill=True, border=1, align='L')
    
    # Fila 1: Nombre Completo
    pdf.set_font('Arial', 'B', SIZE)
    pdf.cell(35, H_CELDA, "Nombre Completo:", border='LTB', align='L')
    pdf.set_font('Arial', '', SIZE)
    pdf.cell(epw-35, H_CELDA, f"{p.get('nombres', '')} {p.get('apellidos', '')}", border='RTB', ln=1, align='L')
    
    # Fila 2: Edad, Sexo, DNI
    w3 = epw / 3
    pdf.set_font('Arial', 'B', SIZE); pdf.cell(20, H_CELDA, "Edad:", border='LTB')
    pdf.set_font('Arial', '', SIZE); pdf.cell(w3-20, H_CELDA, str(p.get('edad', '')), border='RTB')
    pdf.set_font('Arial', 'B', SIZE); pdf.cell(20, H_CELDA, "Sexo:", border='LTB')
    pdf.set_font('Arial', '', SIZE); pdf.cell(w3-20, H_CELDA, str(p.get('sexo', '')), border='RTB')
    pdf.set_font('Arial', 'B', SIZE); pdf.cell(20, H_CELDA, "DNI:", border='LTB')
    pdf.set_font('Arial', '', SIZE); pdf.cell(w3-20, H_CELDA, str(p.get('dni', '')), border='RTB', ln=1)

    # Fila 3: Fecha de Nac, Lugar
    pdf.set_font('Arial', 'B', SIZE); pdf.cell(25, H_CELDA, "F. Nac:", border='LTB')
    pdf.set_font('Arial', '', SIZE); pdf.cell(w3-25, H_CELDA, str(p.get('fecha_nac', '')), border='RTB')
    pdf.set_font('Arial', 'B', SIZE); pdf.cell(25, H_CELDA, "Lugar:", border='LTB')
    pdf.set_font('Arial', '', SIZE); pdf.cell((epw-w3)-25, H_CELDA, str(p.get('lugar', '')), border='RTB', ln=1)

    # Fila 4: Est. Civil, Nº hijos, Grado de Instrucción
    pdf.set_font('Arial', 'B', SIZE); pdf.cell(25, H_CELDA, "Est. Civil:", border='LTB')
    pdf.set_font('Arial', '', SIZE); pdf.cell(w3-25, H_CELDA, str(p.get('estado_civil', '')), border='RTB')
    pdf.set_font('Arial', 'B', SIZE); pdf.cell(25, H_CELDA, "Nº hijos:", border='LTB')
    pdf.set_font('Arial', '', SIZE); pdf.cell(15, H_CELDA, str(p.get('hijos', '')), border='RTB')
    pdf.set_font('Arial', 'B', SIZE); pdf.cell(35, H_CELDA, "Grado Instrucción:", border='LTB')
    pdf.set_font('Arial', '', SIZE); pdf.cell(epw-w3-75, H_CELDA, str(p.get('instruccion', '')), border='RTB', ln=1)

    # Fila 5: Centro de Trabajo, Dirección
    pdf.set_font('Arial', 'B', SIZE); pdf.cell(35, H_CELDA, "Centro Trabajo:", border='LTB')
    pdf.set_font('Arial', '', SIZE); pdf.cell(w3-15, H_CELDA, str(p.get('trabajo', '')), border='RTB')
    pdf.set_font('Arial', 'B', SIZE); pdf.cell(20, H_CELDA, "Dirección:", border='LTB')
    pdf.set_font('Arial', '', SIZE); pdf.cell(epw-w3-40, H_CELDA, str(p.get('direccion', '')), border='RTB', ln=1)

    # Fila 6: Teléfono, Fecha Entrevista
    pdf.set_font('Arial', 'B', SIZE); pdf.cell(25, H_CELDA, "Teléfono:", border='LTB')
    pdf.set_font('Arial', '', SIZE); pdf.cell(w3-5, H_CELDA, str(p.get('telefono', '')), border='RTB')
    pdf.set_font('Arial', 'B', SIZE); pdf.cell(40, H_CELDA, "Fecha Entrevista:", border='LTB')
    pdf.set_font('Arial', '', SIZE); pdf.cell(epw-w3-60, H_CELDA, str(datos.get('fecha_entrevista', '')), border='RTB', ln=1)

   # Fila 7: Modalidad Detallada
    f = datos.get('filiacion', {})
    
    # Definimos los anchos de las celdas previas
    w_mod_lab, w_mod_val = 25, w3 - 10
    w_te_lab, w_te_val = 15, 15
    w_dia_lab, w_dia_val = 30, 20
    w_acc_lab = 20
    
    # Calculamos cuánto espacio queda para el valor de "Acceso"
    # epw es el ancho total disponible entre márgenes
    ancho_usado = w_mod_lab + w_mod_val + w_te_lab + w_te_val + w_dia_lab + w_dia_val + w_acc_lab
    ancho_restante = epw - ancho_usado

    pdf.set_font('Arial', 'B', SIZE); pdf.cell(w_mod_lab, H_CELDA, "Modalidad:", border='LTB')
    pdf.set_font('Arial', '', SIZE); pdf.cell(w_mod_val, H_CELDA, f"{f.get('modalidad', '')} (Turno: {f.get('turno', '')})", border='RTB')
    
    pdf.set_font('Arial', 'B', SIZE); pdf.cell(w_te_lab, H_CELDA, "TE:", border='LTB')
    pdf.set_font('Arial', '', SIZE); pdf.cell(w_te_val, H_CELDA, str(f.get('te', '')), border='RTB')
    
    pdf.set_font('Arial', 'B', SIZE); pdf.cell(w_dia_lab, H_CELDA, "T. Diálisis:", border='LTB')
    pdf.set_font('Arial', '', SIZE); pdf.cell(w_dia_val, H_CELDA, str(f.get('t_dialisis', '')), border='RTB')
    
    pdf.set_font('Arial', 'B', SIZE); pdf.cell(w_acc_lab, H_CELDA, "Acceso:", border='LTB')
    # Usamos ancho_restante para que el cuadro termine exactamente en el borde derecho
    pdf.set_font('Arial', '', SIZE); pdf.cell(ancho_restante, H_CELDA, str(f.get('acceso', '')), border='RTB', ln=1)

    # Bloques de Texto: Antecedentes [cite: 9, 71]
    pdf.set_font('Arial', 'B', SIZE)
    pdf.set_x(10)
    pdf.ln(2)
    pdf.cell(epw, 6, "Antecedentes:", ln=1, align='L')
    pdf.set_font('Arial', '', SIZE)
    pdf.multi_cell(epw, 5, str(datos.get('antecedentes', 'N/R')).strip(), border=1, align='L')
    
    # Bloques de Texto: Conciencia (CORREGIDO ALINEACIÓN) [cite: 10, 72]
    pdf.ln(2)
    pdf.set_font('Arial', 'B', SIZE)
    pdf.set_x(10)
    pdf.ln(2)
    pdf.cell(epw, 6, "Conciencia de la enfermedad:", ln=1, align='L')
    pdf.set_font('Arial', '', SIZE)
    pdf.set_x(10)
    pdf.multi_cell(epw, 5, str(datos.get('conciencia', 'N/R')).strip(), border=1, align='L')
    pdf.ln(S_ESPACIO)

    # II. DATOS COMPLEMENTARIOS [cite: 11, 74]
    pdf.set_font('Arial', 'B', SIZE)
    pdf.set_x(10)
    pdf.ln(2)
    pdf.cell(epw, H_CELDA, 'II. Datos Complementarios', ln=1, fill=True, border=1, align='L')
    
    pdf.cell(epw, 6, "1. Composición Familiar:", ln=1, align='L')
    familia = datos.get('familia', {})
    cols = ["Nombre", "Edad", "Parentesco", "E. Civil", "Localización", "Ocupación", "Relación"]
    w_col = epw / 7
    pdf.set_font('Arial', 'B', 8)
    for col in cols: pdf.cell(w_col, 6, col, border=1, align='C', fill=True)
    pdf.ln()
    
    pdf.set_font('Arial', '', 8)
    n_f = len(familia.get('Nombre', {})) if isinstance(familia, dict) and 'Nombre' in familia else 0
    if n_f > 0:
        for i in familia['Nombre'].keys():
            for col in cols:
                val = str(familia.get(col, {}).get(i, ''))
                pdf.cell(w_col, 6, val, border=1, align='L')
            pdf.ln()
    else: pdf.cell(epw, 6, "No registrado", border=1, ln=1, align='C')

    # Dinámica y Cuidador [cite: 14, 16, 78, 82]
    pdf.ln(2)
    pdf.set_font('Arial', 'B', SIZE)
    pdf.set_x(10)
    pdf.cell(30, H_CELDA, "Dinámica:", border='LTB', align='L')
    pdf.set_font('Arial', '', SIZE)
    pdf.cell(epw*0.5-30, H_CELDA, str(datos.get('dinamica', 'N/R')), border='RTB', align='L')
    
    pdf.set_font('Arial', 'B', SIZE)
    pdf.cell(30, H_CELDA, "Cuidador:", border='LTB', align='L')
    pdf.set_font('Arial', '', SIZE)
    pdf.cell(epw*0.5-30, H_CELDA, str(datos.get('nombre_cuidador', 'N/R')), border='RTB', ln=1, align='L')

    pdf.set_font('Arial', 'B', SIZE)
    pdf.set_x(10)
    pdf.ln(2)
    pdf.cell(epw, 6, "Detalles de Entorno Familiar:", ln=1, align='L')
    pdf.set_font('Arial', '', SIZE)
    pdf.ln(2)
    pdf.multi_cell(epw, 5, f"Dinámica: {datos.get('det_dinamica', 'N/R')}\nRoles: {', '.join(datos.get('rol_cuidador', []))}".strip(), border=1, align='L')
    
    # Educación (CORREGIDO ALINEACIÓN) [cite: 19, 94]
    pdf.ln(2)
    pdf.set_font('Arial', 'B', SIZE)
    pdf.set_x(10)
    pdf.cell(epw, 6, "4. Educación:", ln=1, align='L')
    pdf.set_font('Arial', '', SIZE)
    pdf.set_x(10)
    pdf.ln(2)
    pdf.multi_cell(epw, 5, str(datos.get('educacion', 'N/R')).strip(), border=1, align='L')

    # Experiencia Laboral [cite: 21, 81]
    pdf.ln(2)
    pdf.set_font('Arial', 'B', SIZE)
    pdf.set_x(10)
    pdf.cell(epw, 6, "5. Experiencia Laboral:", ln=1, align='L')
    pdf.set_font('Arial', '', SIZE)
    pdf.set_x(10)
    pdf.ln(2)
    pdf.multi_cell(epw, 5, str(datos.get('exp_laboral', 'N/R')).strip(), border=1, align='L')

    # 6. Adherencia con Gráfico de Barras (CON RESET DE COLOR)
    pdf.ln(2)
    pdf.set_font('Arial', 'B', SIZE)
    pdf.set_x(10)
    pdf.cell(epw, 6, "6. Niveles de Adherencia (Escala 1 al 5):", ln=1, align='L')
    pdf.ln(2) 
    
    adh = datos.get('adherencia', {})
    indicadores = [("Asis", adh.get('asistencia', 0)), ("Dieta", adh.get('dieta', 0)), 
                   ("Farm", adh.get('farma', 0)), ("Hig", adh.get('higiene', 0))]
    
    w_col = epw / 4
    h_barra = 2.5
    y_txt = pdf.get_y()
    
    pdf.set_draw_color(180, 180, 180)
    pdf.rect(10, y_txt, epw, 12) 

    for i, (label, valor) in enumerate(indicadores):
        x_pos = 10 + (i * w_col)
        pdf.set_xy(x_pos + 2, y_txt + 1)
        pdf.set_font('Arial', 'B', 8)
        pdf.cell(w_col, 4, f"{label}: {valor}", align='L')
        
        pdf.set_fill_color(240, 240, 240) # Fondo gris de la barra
        pdf.rect(x_pos + 2, y_txt + 6, w_col - 5, h_barra, style='F')
        
        try:
            ancho_progreso = (float(valor) / 5) * (w_col - 5)
            pdf.set_fill_color(0, 168, 150) # Verde Teal solo para el progreso
            pdf.rect(x_pos + 2, y_txt + 6, ancho_progreso, h_barra, style='F')
        except: pass

    # RESET CRÍTICO: Devolver el color de relleno al PLOMO de los títulos
    pdf.set_fill_color(245, 245, 245) 
    pdf.set_y(y_txt + 12)
    pdf.set_x(10)
    pdf.set_font('Arial', 'B', SIZE)
    pdf.cell(epw*0.4, H_CELDA, f" Hidratación: {datos.get('hidrico', 'N/R')}", border=1, align='L')
    pdf.cell(epw*0.6, H_CELDA, f" Hábitos: {', '.join(datos.get('habitos', []))}", border=1, ln=1, align='L')

    # III. EXAMEN MENTAL [cite: 32, 111]
    if pdf.get_y() + 65 > 270: pdf.add_page()
    pdf.set_font('Arial', 'B', SIZE)
    pdf.set_x(10)
    pdf.cell(epw, H_CELDA, 'III. Examen Mental', ln=1, fill=True, border=1, align='L')
    
    em = datos.get('examen_mental', {})
    y_m = pdf.get_y()
    
    # Columnas: Observaciones y Afecto [cite: 33, 35, 112, 114]
    pdf.set_font('Arial', 'B', SIZE)
    pdf.cell(epw/2, 5, "1. Observaciones:", border='LT', align='L')
    pdf.set_xy(10 + epw/2, y_m)
    pdf.cell(epw/2, 5, "2. Estado afectivo:", border='LT', ln=1, align='L')
    
    y_cont = pdf.get_y()
    pdf.set_font('Arial', '', SIZE)
    pdf.multi_cell(epw/2, 5, str(em.get('obs', 'N/R')).strip(), border='LR')
    y1 = pdf.get_y()
    pdf.set_xy(10 + epw/2, y_cont)
    pdf.multi_cell(epw/2, 5, str(em.get('afe', 'N/R')).strip(), border='LR')
    y2 = pdf.get_y()
    
    # Columnas: Cognoscitivos y Voluntaria [cite: 37, 39, 113, 115]
    pdf.set_y(max(y1, y2))
    pdf.set_x(10)
    pdf.cell(epw, 0, "", border='T', ln=1)
    y_m2 = pdf.get_y()
    pdf.set_font('Arial', 'B', SIZE)
    pdf.cell(epw/2, 5, "3. Cognoscitivos:", border='L', align='L')
    pdf.set_xy(10 + epw/2, y_m2)
    pdf.cell(epw/2, 5, "4. Voluntaria:", border='L', ln=1, align='L')
    
    y_cont2 = pdf.get_y()
    pdf.set_font('Arial', '', SIZE)
    pdf.multi_cell(epw/2, 5, str(em.get('cog', 'N/R')).strip(), border='LB')
    y3 = pdf.get_y()
    pdf.set_xy(10 + epw/2, y_cont2)
    pdf.multi_cell(epw/2, 5, str(em.get('vol', 'N/R')).strip(), border='LB')
    y4 = pdf.get_y()
    
    pdf.set_y(max(y3, y4))
    pdf.set_x(10)
    pdf.ln(S_ESPACIO)

    # IV. EVALUACIÓN Y V. MONITOREO [cite: 41, 43, 120, 121]
    pdf.set_font('Arial', 'B', SIZE)
    pdf.set_x(10)
    pdf.ln(2)
    pdf.cell(epw, H_CELDA, f"IV. Evaluación (Tests): {', '.join(datos.get('eval_p', []))}", border=1, ln=1, align='L')
    pdf.ln(S_ESPACIO)

    if pdf.get_y() + 90 > 270: pdf.add_page()
    pdf.set_x(10)
    pdf.ln(2)
    pdf.cell(epw, H_CELDA, 'V. Monitoreo Bio-conductual (S1-S9)', ln=1, fill=True, border=1, align='L')
    if grafico_path:
        w_g = epw * 0.70
        pdf.image(grafico_path, x=(pdf.w - w_g)/2, y=pdf.get_y() + 2, w=w_g)
        pdf.ln(75) 
    pdf.ln(S_ESPACIO)

    # VI y VII. DIAGNÓSTICO E INTERVENCIÓN [cite: 45, 46, 145, 149]
    if pdf.get_y() + 45 > 270: pdf.add_page()
    pdf.set_font('Arial', 'B', SIZE)
    pdf.set_x(10)
    pdf.ln(20)
    pdf.cell(epw, H_CELDA, 'VI. Diagnóstico y VII. Intervención', ln=1, fill=True, border=1, align='L')
    
    pdf.set_x(10)
    pdf.cell(epw, 6, "Diagnóstico:", ln=1, align='L')
    pdf.set_font('Arial', '', SIZE)
    pdf.set_x(10)
    pdf.multi_cell(epw, 5, str(datos.get('diagnostico', 'N/R')).strip(), border=1, align='L')
    
    pdf.set_font('Arial', 'B', SIZE)
    pdf.set_x(10)
    pdf.cell(epw, 6, f"Intervención ({datos.get('intervencion', 'N/R')}):", ln=1, align='L')
    pdf.set_font('Arial', '', SIZE)
    pdf.set_x(10)
    pdf.multi_cell(epw, 5, str(datos.get('det_interv', 'Sin detalles.')).strip(), border=1, align='L')

    return bytes(pdf.output())

def generar_reporte_administrativo_limpio(fecha, datos_jornada):
    # Usamos tu clase FichaPDF para mantener logo y encabezado
    pdf = FichaPDF() 
    pdf.add_page()
    epw = pdf.epw
    
    pdf.set_font('Arial', 'B', 12)
    pdf.cell(epw, 10, f'REGISTRO DIARIO DE PRODUCCIÓN ASISTENCIAL - {fecha}', ln=1, align='C', border='B')
    pdf.ln(5)
    
    for act, pacientes in datos_jornada.items():
        # Encabezado de actividad
        pdf.set_fill_color(240, 240, 240)
        pdf.set_font('Arial', 'B', 10)
        pdf.cell(epw, 8, f" ACTIVIDAD: {act.upper()}", ln=1, fill=True, border=1)
        
        # Cabecera tabla
        pdf.set_font('Arial', 'B', 9)
        pdf.cell(epw*0.25, 7, " DNI", border=1, align='C')
        pdf.cell(epw*0.75, 7, " APELLIDOS Y NOMBRES", border=1, ln=1, align='C')
        
        pdf.set_font('Arial', '', 9)
        if pacientes:
            for p in pacientes:
                pdf.cell(epw*0.25, 7, f" {p['dni']}", border=1)
                pdf.cell(epw*0.75, 7, f" {p['nombre']}", border=1, ln=1)
            
            # Contador pequeño al final de cada sección
            pdf.set_font('Arial', 'I', 8)
            pdf.cell(epw, 6, f" Total pacientes: {len(pacientes)}", ln=1, align='R')
        else:
            pdf.cell(epw, 7, " Sin pacientes registrados.", border=1, ln=1, align='C')
        
        pdf.ln(4)

    return bytes(pdf.output())

def generar_reporte_administrativo_limpio(fecha, datos_jornada):
    """Genera el PDF administrativo con lista de pacientes (Solo DNI y Nombres)."""
    pdf = FichaPDF() 
    pdf.add_page()
    epw = pdf.epw
    
    pdf.set_font('Arial', 'B', 12)
    pdf.cell(epw, 10, f'REGISTRO DIARIO DE PRODUCCIÓN ASISTENCIAL - {fecha}', ln=1, align='C', border='B')
    pdf.ln(5)
    
    for act, pacientes in datos_jornada.items():
        pdf.set_fill_color(240, 240, 240)
        pdf.set_font('Arial', 'B', 10)
        pdf.cell(epw, 8, f" ACTIVIDAD: {act.upper()}", ln=1, fill=True, border=1)
        
        pdf.set_font('Arial', 'B', 9)
        pdf.cell(epw*0.25, 7, " DNI", border=1, align='C')
        pdf.cell(epw*0.75, 7, " APELLIDOS Y NOMBRES", border=1, ln=1, align='C')
        
        pdf.set_font('Arial', '', 9)
        if pacientes:
            for p in pacientes:
                pdf.cell(epw*0.25, 7, f" {p['dni']}", border=1)
                pdf.cell(epw*0.75, 7, f" {p['nombre']}", border=1, ln=1)
            
            pdf.set_font('Arial', 'I', 8)
            pdf.cell(epw, 6, f" Total pacientes: {len(pacientes)}", ln=1, align='R')
        else:
            pdf.cell(epw, 7, " Sin pacientes registrados.", border=1, ln=1, align='C')
        pdf.ln(4)

    return bytes(pdf.output())

def generar_word_evoluciones(fecha, datos_jornada):
    """Genera el documento Word con las evoluciones redactadas para copiar y editar."""
    doc = Document()
    
    # Configuración de fuente estándar
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # Título
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo.add_run(f"REGISTRO CLÍNICO DE EVOLUCIONES - {fecha}")
    run.bold = True
    run.size = Pt(14)

    for act, pacientes in datos_jornada.items():
        if pacientes:
            # Encabezado de Actividad
            doc.add_paragraph("\n")
            h_act = doc.add_paragraph()
            run_act = h_act.add_run(f"ACTIVIDAD: {act.upper()}")
            run_act.bold = True
            run_act.underline = True
            
            for p in pacientes:
                # Datos del Paciente
                p_para = doc.add_paragraph()
                p_para.add_run(f"PACIENTE: ").bold = True
                p_para.add_run(f"{p['nombre']} (DNI: {p['dni']})")
                
                # Evolución
                evol_para = doc.add_paragraph()
                evol_para.add_run("EVOLUCIÓN: ").bold = True
                evol_para.add_run(p['texto'] if p['texto'] else "Sin redacción registrada.")
                
                # Línea separadora sutil
                doc.add_paragraph("-----------------------------------------------------------------")
            
            doc.add_page_break() # Una página por cada tipo de actividad

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def generar_informe_mensual_secuencial(mes, anio, df):
    """
    Genera un PDF organizado por cuadros independientes de actividad,
    con márgenes corregidos y alineación centrada.
    """
    pdf = FichaPDF()
    pdf.set_margins(left=10, top=10, right=10) # Forzamos márgenes simétricos
    pdf.add_page()
    epw = pdf.epw # Ancho útil (aprox 190mm en A4)
    
    # --- ENCABEZADO (Corregido para centrar en página) ---
    pdf.set_font('Arial', 'B', 14)
    pdf.set_text_color(0, 0, 0)
    # Usamos ancho 0 para que ocupe toda la página y align='C' para centrar real
    pdf.cell(0, 10, 'REGISTRO MENSUAL DE PRODUCCIÓN ASISTENCIAL', ln=1, align='C')
    
    pdf.set_font('Arial', 'B', 10)
    pdf.cell(0, 7, f'Periodo: {mes}/{anio} - Servicio de Psicología', ln=1, align='C')
    pdf.cell(0, 7, f'Ps. Magaly Gonzales Aliaga', ln=1, align='C')
    
    # Línea decorativa horizontal
    pdf.line(10, pdf.get_y(), 200, pdf.get_y())
    pdf.ln(5)

    # --- I. RESUMEN ESTADÍSTICO ---
    pdf.set_fill_color(230, 230, 230)
    pdf.set_font('Arial', 'B', 10)
    pdf.set_x(10) # Aseguramos inicio en el margen izquierdo
    pdf.cell(epw, 8, " I. RESUMEN DE PRODUCTIVIDAD MENSUAL", ln=1, fill=True, border=1)
    
    stats = df['tipo_cita'].value_counts()
    tipos_orden = ["Atención Psicológica", "Intervención Individual", "Intervención Familiar", "Evaluación Diagnóstica"]
    
    pdf.set_font('Arial', 'B', 9)
    pdf.set_x(10)
    pdf.cell(epw*0.7, 7, " ACTIVIDAD / TIPO DE ATENCIÓN", border=1, align='C')
    pdf.cell(epw*0.3, 7, " TOTAL", border=1, ln=1, align='C')
    
    pdf.set_font('Arial', '', 9)
    for tipo in tipos_orden:
        cantidad = stats.get(tipo, 0)
        pdf.set_x(10)
        pdf.cell(epw*0.7, 7, f" {tipo}", border=1)
        pdf.cell(epw*0.3, 7, f" {cantidad}", border=1, ln=1, align='C')
    
    pdf.set_x(10)
    pdf.set_fill_color(245, 245, 245)
    pdf.set_font('Arial', 'B', 9)
    pdf.cell(epw*0.7, 8, " TOTAL GENERAL DE ATENCIONES", border=1, fill=True)
    pdf.cell(epw*0.3, 8, f" {len(df)}", border=1, ln=1, align='C', fill=True)
    pdf.ln(10)

    # --- II. DESGLOSE CRONOLÓGICO POR CUADROS ---
    for tipo_act in tipos_orden:
        df_tipo = df[df['tipo_cita'] == tipo_act].copy()
        if not df_tipo.empty:
            df_tipo['fecha_dt'] = pd.to_datetime(df_tipo['fecha'])
            df_tipo = df_tipo.sort_values(by='fecha_dt')
            
            # Título del Cuadro
            pdf.set_x(10)
            pdf.set_fill_color(240, 240, 240); pdf.set_font('Arial', 'B', 10)
            pdf.cell(epw, 9, f" CUADRO: {tipo_act.upper()}", ln=1, fill=True, border=1)
            
            # Cabecera de Tabla
            pdf.set_x(10)
            pdf.set_font('Arial', 'B', 8)
            pdf.cell(epw*0.12, 8, " FECHA", border=1, align='C')
            pdf.cell(epw*0.30, 8, " APELLIDOS Y NOMBRES", border=1, align='C')
            pdf.cell(epw*0.10, 8, " DNI", border=1, align='C')
            pdf.cell(epw*0.04, 8, " E", border=1, align='C')
            pdf.cell(epw*0.04, 8, " S", border=1, align='C')
            pdf.cell(epw*0.40, 8, " DIAGNÓSTICO / EVALUACIÓN", border=1, ln=1, align='C')

            fechas_en_tipo = df_tipo['fecha_dt'].dt.strftime('%Y-%m-%d').unique()
            
            for f_str in fechas_en_tipo:
                sub_f = df_tipo[df_tipo['fecha_dt'].dt.strftime('%Y-%m-%d') == f_str]
                num_reg = len(sub_f)
                h_fila = 7
                h_bloque = num_reg * h_fila

                # Gestión de saltos de página
                if pdf.get_y() + h_bloque > 270:
                    pdf.add_page()
                    pdf.set_font('Arial', 'B', 8)
                    pdf.set_x(10)
                    pdf.cell(epw*0.12, 8, " FECHA", border=1, align='C')
                    pdf.cell(epw*0.30, 8, " APELLIDOS Y NOMBRES", border=1, align='C')
                    pdf.cell(epw*0.10, 8, " DNI", border=1, align='C')
                    pdf.cell(epw*0.04, 8, " E", border=1, align='C')
                    pdf.cell(epw*0.04, 8, " S", border=1, align='C')
                    pdf.cell(epw*0.40, 8, " DIAGNÓSTICO / EVALUACIÓN", border=1, ln=1, align='C')

                x_i, y_i = 10, pdf.get_y() # Forzamos x_i a 10 (margen izquierdo)

                # --- BLOQUE DE FECHA ---
                pdf.set_xy(x_i, y_i)
                pdf.cell(epw*0.12, h_bloque, "", border=1)
                
                pos_y_central = y_i + (h_bloque / 2) - (h_fila / 2)
                pdf.set_xy(x_i, pos_y_central)
                pdf.set_font('Arial', 'B', 7)
                pdf.cell(epw*0.12, h_fila, str(f_str), align='C')
                
                # Volver para llenar datos del paciente
                pdf.set_xy(x_i + (epw*0.12), y_i)

                for idx, (_, row) in enumerate(sub_f.iterrows()):
                    pdf.set_font('Arial', '', 7)
                    curr_y = pdf.get_y()
                    
                    pdf.cell(epw*0.30, h_fila, f" {row['apellidos']} {row['nombres']}"[:35], border=1)
                    pdf.cell(epw*0.10, h_fila, str(row['dni']), border=1, align='C')
                    pdf.cell(epw*0.04, h_fila, str(row['edad']), border=1, align='C')
                    pdf.cell(epw*0.04, h_fila, str(row['sexo']), border=1, align='C')
                    
                    # Diagnóstico
                    x_pos_diag = pdf.get_x()
                    pdf.multi_cell(epw*0.40, h_fila, str(row['diagnostico']) if row['diagnostico'] else "N/R", border=1)
                    
                    # Resetear X para la siguiente fila de la misma fecha
                    if idx < num_reg - 1:
                        pdf.set_xy(x_i + (epw*0.12), curr_y + h_fila)
            
            pdf.ln(8) 

    return bytes(pdf.output())